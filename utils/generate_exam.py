import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Set cell padding (in twentieths of a point, dxas)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn(f'w:{key}'), str(val))
            tcBorders.append(element)
        else:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'none')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def build_exam():
    doc = Document()
    
    # Standard margins (0.75 in / ~1.9 cm) to balance space over exactly 2 pages nicely
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles Setup (Arial 10pt for readability)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(10)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33)

    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Arial'
    font_h1.size = Pt(11.5)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0x04, 0x78, 0x57)

    # 1. Header Table (Logo + Info)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(2.0)
    header_table.columns[1].width = Inches(5.0)
    
    cell_logo = header_table.rows[0].cells[0]
    logo_path = r"c:\Users\HP 250 G10\Documents\GITHUT\Centuria\SOCIOLOGIA - JUNIO\Asistencia\public\logo_centuria.png"
    if os.path.exists(logo_path):
        run_logo = cell_logo.paragraphs[0].add_run()
        run_logo.add_picture(logo_path, width=Inches(1.8))
        cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        cell_logo.paragraphs[0].text = "[Logo Centuria]"
        cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell_info = header_table.rows[0].cells[1]
    p_info = cell_info.paragraphs[0]
    p_info.paragraph_format.line_spacing = 1.15
    p_info.paragraph_format.space_before = Pt(0)
    p_info.paragraph_format.space_after = Pt(0)
    
    run_inst = p_info.add_run("INSTITUTO SUPERIOR DE CIENCIAS EMPRESARIALES CENTURIA\n")
    run_inst.bold = True
    run_inst.font.size = Pt(10)
    run_inst.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
    
    run_det = p_info.add_run(
        "Materia: Sociología  |  Profesor: C.P Mg. Juan Carlos Montiel\n"
        "Tipo de Evaluación: Cuestionario Evaluativo de Criterio\n"
    )
    run_det.font.size = Pt(9.0)
    
    set_cell_borders(cell_logo, right={'sz': 4, 'val': 'single', 'color': 'D3D3D3'})
    set_cell_borders(cell_info)
    set_cell_margins(cell_logo, top=40, bottom=40, left=30, right=60)
    set_cell_margins(cell_info, top=40, bottom=40, left=60, right=30)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(6)
    p_spacer.paragraph_format.space_after = Pt(6)
    
    # Student Metadata Box
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(3.2)
    meta_table.columns[1].width = Inches(1.9)
    meta_table.columns[2].width = Inches(1.9)
    
    c0 = meta_table.rows[0].cells[0]
    c1 = meta_table.rows[0].cells[1]
    c2 = meta_table.rows[0].cells[2]
    
    p_stud = c0.paragraphs[0]
    p_stud.paragraph_format.space_after = Pt(0)
    r_stud_lbl = p_stud.add_run("Estudiante: ")
    r_stud_lbl.bold = True
    r_stud_lbl.font.size = Pt(9.0)
    r_stud_val = p_stud.add_run("______________________________")
    r_stud_val.font.size = Pt(9.0)
    
    p_date = c1.paragraphs[0]
    p_date.paragraph_format.space_after = Pt(0)
    r_date_lbl = p_date.add_run("Fecha: ")
    r_date_lbl.bold = True
    r_date_lbl.font.size = Pt(9.0)
    r_date_val = p_date.add_run("____/____/2026")
    r_date_val.font.size = Pt(9.0)
    
    p_score = c2.paragraphs[0]
    p_score.paragraph_format.space_after = Pt(0)
    r_score_lbl = p_score.add_run("Puntaje: ")
    r_score_lbl.bold = True
    r_score_lbl.font.size = Pt(9.0)
    r_score_val = p_score.add_run("_______ / 25 pts.")
    r_score_val.font.size = Pt(9.0)
    
    border_style = {'sz': 4, 'val': 'single', 'color': 'D3D3D3'}
    for idx, cell in enumerate(meta_table.rows[0].cells):
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        set_cell_shading(cell, 'F9F9F9')
        set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
            
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("EVALUACIÓN ESCRITA DE SOCIOLOGÍA")
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
    
    # -------------------------------------------------------------
    # SECCIÓN I: ÍTEMS DE SELECCIÓN MÚLTIPLE
    # -------------------------------------------------------------
    h1_sec1 = doc.add_paragraph(style='Heading 1')
    r_sec1 = h1_sec1.add_run("I. SELECCIÓN MÚLTIPLE (10 Puntos - 1 pt. c/u)")
    h1_sec1.paragraph_format.space_before = Pt(6)
    h1_sec1.paragraph_format.space_after = Pt(4)
    
    mc_questions = [
        ("1. ¿Quién propuso la “Ley de los Tres Estados”?", ["a) É. Durkheim", "b) M. Weber", "c) A. Comte", "d) K. Marx"]),
        ("2. Según Comte, el estado teológico se caracteriza por:", ["a) Experimentación científica", "b) Explicaciones por fuerzas divinas", "c) Causas económicas", "d) Burocracia moderna"]),
        ("3. En el estado positivo, el conocimiento se basa en:", ["a) Revelación divina", "b) Especulación abstracta", "c) Observación y método científico", "d) Costumbres tradicionales"]),
        ("4. Para Durkheim, un “hecho social” es:", ["a) Un sentimiento individual", "b) Una norma externa que ejerce coerción", "c) Una preferencia personal", "d) Una opinión subjetiva"]),
        ("5. La solidaridad mecánica se vincula con sociedades:", ["a) Industriales y modernas", "b) Simples y tradicionales homogéneas", "c) Sin normas sociales", "d) Basadas solo en contratos"]),
        ("6. La solidaridad orgánica se relaciona principalmente con:", ["a) Poca especialización laboral", "b) Diferenciación de funciones e interdependencia", "c) Ausencia de normas", "d) Sociedades tribales"]),
        ("7. Para Weber, la “acción social” es:", ["a) Un acto biológico reflejo", "b) Una conducta con sentido subjetivo orientada a otros", "c) Un movimiento instintivo", "d) Un hecho puramente natural"]),
        ("8. La racionalidad formal en Weber se observa típicamente en:", ["a) Rituales tradicionales", "b) Organizaciones burocráticas reglamentadas", "c) Relaciones afectivas", "d) Actos artísticos libres"]),
        ("9. Para Marx, la estructura económica de la sociedad:", ["a) No influye en las leyes", "b) Determina la superestructura política y cultural", "c) Es totalmente independiente", "d) Solo influye en el agro"]),
        ("10. El concepto de “anomia” en Durkheim se refiere a:", ["a) Ausencia de leyes físicas", "b) Debilidad o falta de normas sociales claras", "c) Exceso de moral religiosa", "d) Orden social estricto"])
    ]
    
    for q_text, opts in mc_questions:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(3)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(q_text)
        r_q.bold = True
        
        p_o = doc.add_paragraph()
        p_o.paragraph_format.left_indent = Inches(0.25)
        p_o.paragraph_format.space_after = Pt(2)
        p_o.add_run(f"{opts[0]}      {opts[1]}      {opts[2]}      {opts[3]}")
        
    # -------------------------------------------------------------
    # SECCIÓN II: ÍTEM DE COMPARACIÓN
    # -------------------------------------------------------------
    h1_sec2 = doc.add_paragraph(style='Heading 1')
    r_sec2 = h1_sec2.add_run("II. CUADRO COMPARATIVO (10 Puntos)")
    h1_sec2.paragraph_format.space_before = Pt(12)
    h1_sec2.paragraph_format.space_after = Pt(6)
    
    p_inst = doc.add_paragraph()
    p_inst.paragraph_format.space_after = Pt(4)
    r_inst = p_inst.add_run("Instrucciones: Complete el cuadro comparativo señalando dos semejanzas, dos diferencias y un ejemplo práctico actual para cada tipo de solidaridad de Émile Durkheim.")
    r_inst.font.italic = True
    r_inst.font.size = Pt(8.5)

    comp_table = doc.add_table(rows=4, cols=3)
    comp_table.autofit = False
    comp_table.columns[0].width = Inches(1.8)
    comp_table.columns[1].width = Inches(2.6)
    comp_table.columns[2].width = Inches(2.6)
    
    headers = ["Criterio", "Solidaridad Mecánica", "Solidaridad Orgánica"]
    for i, title in enumerate(headers):
        cell = comp_table.rows[0].cells[i]
        cell.paragraphs[0].text = title
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.0)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "047857")
        set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
        
    rows_data = [
        ("Semejanzas (Indicar 2)", "", ""),
        ("Diferencias (Indicar 2)", "", ""),
        ("Ejemplo Práctico Actual", "", "")
    ]
    
    border_cell = {'sz': 4, 'val': 'single', 'color': 'D3D3D3'}
    for row_idx, (crit, col1, col2) in enumerate(rows_data, start=1):
        row = comp_table.rows[row_idx]
        c0 = row.cells[0]
        c0.paragraphs[0].text = crit
        c0.paragraphs[0].runs[0].font.size = Pt(8.5)
        c0.paragraphs[0].runs[0].bold = True
        set_cell_shading(c0, "F2F2F2")
        set_cell_margins(c0, top=120, bottom=120, left=60, right=60)
        set_cell_borders(c0, top=border_cell, bottom=border_cell, left=border_cell, right=border_cell)
        
        for c_idx in (1, 2):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.add_run("\n\n\n\n")  # Generous space for student writing
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            set_cell_borders(cell, top=border_cell, bottom=border_cell, left=border_cell, right=border_cell)

    # -------------------------------------------------------------
    # SECCIÓN III: ÍTEMS DE VERDADERO O FALSO
    # -------------------------------------------------------------
    h1_sec3 = doc.add_paragraph(style='Heading 1')
    r_sec3 = h1_sec3.add_run("III. VERDADERO O FALSO CON COMENTARIO BREVE (5 Puntos - 1 pt. c/u)")
    h1_sec3.paragraph_format.space_before = Pt(12)
    h1_sec3.paragraph_format.space_after = Pt(6)
    
    vf_items = [
        "1. (   ) Para Comte, el estado positivo exige abandonar las explicaciones místicas y basarse en la ciencia.",
        "2. (   ) Para Durkheim, los hechos sociales son externos y ejercen fuerza colectiva coercitiva sobre el individuo.",
        "3. (   ) Según Weber, la acción social siempre se orienta por reglamentos escritos y formales.",
        "4. (   ) Para Marx, la lucha de clases y los factores económicos son el motor del cambio histórico.",
        "5. (   ) La anomia es un estado donde las reglas sociales pierden fuerza o se vuelven confusas."
    ]
    
    for item in vf_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_before = Pt(3)
        p_item.paragraph_format.space_after = Pt(2)
        r_item = p_item.add_run(item)
        r_item.bold = True
        r_item.font.size = Pt(9.5)
        
        p_com = doc.add_paragraph()
        p_com.paragraph_format.left_indent = Inches(0.3)
        p_com.paragraph_format.space_after = Pt(6)
        r_lbl = p_com.add_run("Comentario: ")
        r_lbl.italic = True
        r_lbl.font.size = Pt(9.0)
        r_line = p_com.add_run("__________________________________________________________________________________")
        r_line.font.size = Pt(9.0)
        
    # Save file
    output_path = r"c:\Users\HP 250 G10\Documents\GITHUT\Centuria\SOCIOLOGIA - JUNIO\Asistencia\Examen_Sociologia_25pts.docx"
    doc.save(output_path)
    print(f"Exam updated successfully at: {output_path}")

if __name__ == "__main__":
    build_exam()
